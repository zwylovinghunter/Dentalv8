"""Polished PDF and Word exporters for the generated Markdown reports.

The application intentionally keeps Markdown as the source of truth.  This
module turns the small, predictable Markdown subset used by the report center
into publication-ready PDF and DOCX documents without changing the detection
data or report semantics.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from hashlib import sha256
from html import escape as html_escape
from pathlib import Path
import re
from typing import Iterable, Sequence
import zipfile

from PIL import Image as PILImage


NAVY = "#123B55"
NAVY_DARK = "#0B263A"
TEAL = "#0F9D8A"
CYAN = "#2AA7C8"
INK = "#243B53"
MUTED = "#66788A"
LINE = "#D8E4EC"
PALE = "#F4F8FB"
PALE_TEAL = "#EAF8F5"
PALE_BLUE = "#EAF3FA"
PALE_AMBER = "#FFF6DF"
PALE_RED = "#FDECEC"
WHITE = "#FFFFFF"


@dataclass
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    rows: list[list[str]] = field(default_factory=list)
    items: list[str] = field(default_factory=list)
    ordered: bool = False
    image_path: Path | None = None
    alt: str = ""


def _clean_inline_text(value: str) -> str:
    text = str(value or "")
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return re.sub(r"\s+", " ", text).strip()


def _split_table_row(line: str) -> list[str]:
    raw = line.strip().strip("|")
    return [_clean_inline_text(cell.replace("\\|", "|")) for cell in re.split(r"(?<!\\)\|", raw)]


def _is_table_separator(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def _parse_image(line: str, base_dir: Path) -> tuple[str, Path] | None:
    match = re.fullmatch(r"\s*!\[([^\]]*)\]\(([^)]+)\)\s*", line)
    if not match:
        return None
    raw_path = match.group(2).strip()
    if raw_path.startswith("data:image/"):
        return None
    image_path = Path(raw_path)
    if not image_path.is_absolute():
        image_path = (base_dir / image_path).resolve()
    return match.group(1).strip(), image_path


def parse_report_markdown(markdown: str, base_dir: str | Path) -> list[MarkdownBlock]:
    """Parse the Markdown subset emitted by the application into layout blocks."""
    lines = str(markdown or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    base = Path(base_dir)
    blocks: list[MarkdownBlock] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if re.fullmatch(r"-{3,}|\*{3,}", stripped):
            blocks.append(MarkdownBlock(kind="rule"))
            index += 1
            continue
        image = _parse_image(line, base)
        if image:
            alt, image_path = image
            blocks.append(MarkdownBlock(kind="image", alt=alt, image_path=image_path))
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if heading:
            blocks.append(MarkdownBlock(kind="heading", level=len(heading.group(1)), text=heading.group(2)))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            rows = [_split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_split_table_row(lines[index]))
                index += 1
            width = max((len(row) for row in rows), default=0)
            rows = [row + [""] * (width - len(row)) for row in rows]
            blocks.append(MarkdownBlock(kind="table", rows=rows))
            continue
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            blocks.append(MarkdownBlock(kind="quote", text="\n".join(quote_lines)))
            continue
        list_match = re.match(r"^\s*(?:([-+*])|(\d+)[.)])\s+(.+)$", line)
        if list_match:
            ordered = bool(list_match.group(2))
            items: list[str] = []
            while index < len(lines):
                current = re.match(r"^\s*(?:([-+*])|(\d+)[.)])\s+(.+)$", lines[index])
                if not current or bool(current.group(2)) != ordered:
                    break
                items.append(current.group(3).strip())
                index += 1
            blocks.append(MarkdownBlock(kind="list", items=items, ordered=ordered))
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                break
            if (
                re.match(r"^#{1,6}\s+", candidate_stripped)
                or candidate_stripped.startswith((">", "|"))
                or re.match(r"^\s*(?:[-+*]|\d+[.)])\s+", candidate)
                or _parse_image(candidate, base)
                or re.fullmatch(r"-{3,}|\*{3,}", candidate_stripped)
            ):
                break
            paragraph_lines.append(candidate_stripped)
            index += 1
        blocks.append(MarkdownBlock(kind="paragraph", text=" ".join(paragraph_lines)))
    return blocks


def export_markdown_bundle(markdown: str, path: str | Path, base_dir: str | Path) -> str:
    """Package Markdown and all referenced local images into one portable ZIP."""
    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    base = Path(base_dir)
    rewritten: list[str] = []
    archived_paths: dict[Path, str] = {}
    asset_index = 0

    for line in str(markdown or "").splitlines():
        image = _parse_image(line, base)
        if image is None:
            rewritten.append(line)
            continue
        alt, image_path = image
        try:
            resolved = image_path.resolve(strict=True)
        except OSError:
            rewritten.append(line)
            continue
        archive_name = archived_paths.get(resolved)
        if archive_name is None:
            asset_index += 1
            safe_stem = re.sub(r"[^0-9A-Za-z._-]+", "_", resolved.stem).strip("._") or "image"
            suffix = resolved.suffix.lower() if resolved.suffix else ".png"
            archive_name = f"assets/{asset_index:02d}_{safe_stem[:72]}{suffix}"
            archived_paths[resolved] = archive_name
        safe_alt = alt.replace("\\", "\\\\").replace("]", "\\]")
        rewritten.append(f"![{safe_alt}]({archive_name})")

    markdown_name = f"{output.stem}.md"
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        archive.writestr(markdown_name, "\n".join(rewritten).encode("utf-8"))
        archive.writestr(
            "README.txt",
            (
                "Open the Markdown file in this folder and keep the assets folder beside it.\r\n"
                "请解压整个压缩包，并保持 Markdown 文件与 assets 文件夹的相对位置不变。\r\n"
            ).encode("utf-8-sig"),
        )
        for source, archive_name in archived_paths.items():
            archive.write(source, archive_name)
    return str(output)


def _extract_cover(blocks: Sequence[MarkdownBlock]) -> tuple[str, list[tuple[str, str]], list[MarkdownBlock]]:
    title = "检测辅助分析报告"
    cover_rows: list[tuple[str, str]] = []
    consumed: set[int] = set()
    for idx, block in enumerate(blocks):
        if block.kind == "heading" and block.level == 1:
            title = _clean_inline_text(block.text)
            consumed.add(idx)
            break
    for idx, block in enumerate(blocks):
        if block.kind != "heading" or block.level > 2:
            continue
        normalized = _clean_inline_text(block.text).lower()
        if normalized not in {"报告封面", "cover"}:
            continue
        consumed.add(idx)
        if idx + 1 < len(blocks) and blocks[idx + 1].kind == "table":
            table_rows = blocks[idx + 1].rows
            for row in table_rows[1:]:
                if len(row) >= 2 and (row[0] or row[1]):
                    cover_rows.append((_clean_inline_text(row[0]), _clean_inline_text(row[1])))
            consumed.add(idx + 1)
        break
    return title, cover_rows, [block for idx, block in enumerate(blocks) if idx not in consumed]


def _metadata_value(rows: Sequence[tuple[str, str]], names: Iterable[str], fallback: str = "-") -> str:
    wanted = {name.lower() for name in names}
    for key, value in rows:
        if key.strip().lower() in wanted:
            return value or fallback
    return fallback


def _is_english(rows: Sequence[tuple[str, str]], title: str) -> bool:
    return any(key.lower() in {"report type", "generated at", "project"} for key, _ in rows) or bool(
        re.search(r"\b(report|dental|candidate)\b", title, re.IGNORECASE)
    )


def _inline_tokens(text: str) -> list[tuple[str, bool, bool]]:
    """Return (text, bold, code) spans for the tiny inline Markdown subset."""
    tokens: list[tuple[str, bool, bool]] = []
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")
    position = 0
    for match in pattern.finditer(str(text or "")):
        if match.start() > position:
            tokens.append((text[position : match.start()], False, False))
        token = match.group(0)
        if token.startswith("**"):
            tokens.append((token[2:-2], True, False))
        elif token.startswith("`"):
            tokens.append((token[1:-1], False, True))
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            tokens.append(((link.group(1) if link else token), False, False))
        position = match.end()
    if position < len(text):
        tokens.append((text[position:], False, False))
    return tokens or [(str(text or ""), False, False)]


def _reportlab_inline(text: str) -> str:
    parts: list[str] = []
    for value, bold, code in _inline_tokens(text):
        safe = html_escape(value).replace("\n", "<br/>")
        if code:
            safe = f"<font color='{TEAL}'>{safe}</font>"
        if bold:
            safe = f"<b>{safe}</b>"
        parts.append(safe)
    return "".join(parts)


def _text_weight(text: str) -> float:
    value = _clean_inline_text(text)
    if not value:
        return 1.0
    weight = sum(1.8 if ord(char) > 127 else 1.0 for char in value)
    return max(1.0, min(weight, 34.0))


def _column_widths(rows: Sequence[Sequence[str]], total: float, min_width: float) -> list[float]:
    if not rows:
        return [total]
    columns = max(len(row) for row in rows)
    weights: list[float] = []
    for col in range(columns):
        values = [row[col] if col < len(row) else "" for row in rows]
        max_weight = max((_text_weight(value) for value in values), default=1.0)
        numeric = all(re.fullmatch(r"[-+]?\d+(?:\.\d+)?%?|[-–—]", _clean_inline_text(value)) for value in values[1:] if value)
        weights.append(min(max_weight, 10.0 if numeric else 26.0))
    available = max(total - min_width * columns, 0)
    weight_sum = sum(weights) or 1.0
    widths = [min_width + available * weight / weight_sum for weight in weights]
    correction = total - sum(widths)
    widths[-1] += correction
    return widths


def _register_reportlab_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    regular_name = "DentalSans"
    bold_name = "DentalSansBold"
    regular_candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf"),
        Path("/usr/share/fonts/opentype/adobe-source-han-sans/SourceHanSansSC-Regular.otf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    bold_candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Bold.otf"),
        Path("/usr/share/fonts/opentype/adobe-source-han-sans/SourceHanSansSC-Bold.otf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]
    regular_path = next((path for path in regular_candidates if path.exists()), None)
    if regular_path is None:
        return "Helvetica", "Helvetica-Bold"
    registered = set(pdfmetrics.getRegisteredFontNames())
    if regular_name not in registered:
        try:
            pdfmetrics.registerFont(TTFont(regular_name, str(regular_path), subfontIndex=0))
        except TypeError:
            pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
    if bold_name not in registered:
        bold_path = next((path for path in bold_candidates if path.exists()), regular_path)
        try:
            pdfmetrics.registerFont(TTFont(bold_name, str(bold_path), subfontIndex=0))
        except TypeError:
            pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
    pdfmetrics.registerFontFamily(regular_name, normal=regular_name, bold=bold_name, italic=regular_name, boldItalic=bold_name)
    return regular_name, bold_name


def export_pdf_from_markdown(markdown: str, path: str | Path, base_dir: str | Path) -> str:
    """Create an A4 PDF with cover, tables, figures, and running page furniture."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        Image,
        LongTable,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    blocks = parse_report_markdown(markdown, base_dir)
    title, cover_rows, body_blocks = _extract_cover(blocks)
    english = _is_english(cover_rows, title)
    font_name, bold_name = _register_reportlab_fonts()
    page_width, page_height = A4
    margin_x = 18 * mm
    available_width = page_width - margin_x * 2

    styles = {
        "cover_kicker": ParagraphStyle(
            "DentalCoverKicker", fontName=bold_name, fontSize=9.5, leading=13, textColor=colors.HexColor(WHITE),
            spaceAfter=6, alignment=TA_LEFT, wordWrap="CJK",
        ),
        "cover_title": ParagraphStyle(
            "DentalCoverTitle", fontName=bold_name, fontSize=25, leading=34, textColor=colors.HexColor(WHITE),
            spaceAfter=7, alignment=TA_LEFT, wordWrap="CJK",
        ),
        "cover_subtitle": ParagraphStyle(
            "DentalCoverSubtitle", fontName=font_name, fontSize=11.5, leading=18, textColor=colors.HexColor("#D9F4F0"),
            alignment=TA_LEFT, wordWrap="CJK",
        ),
        "cover_label": ParagraphStyle(
            "DentalCoverLabel", fontName=font_name, fontSize=7.5, leading=10, textColor=colors.HexColor(MUTED),
            spaceAfter=3, wordWrap="CJK",
        ),
        "cover_value": ParagraphStyle(
            "DentalCoverValue", fontName=bold_name, fontSize=10.5, leading=15, textColor=colors.HexColor(INK),
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "DentalH1", fontName=bold_name, fontSize=17, leading=24, textColor=colors.HexColor(NAVY),
            spaceBefore=16, spaceAfter=8, keepWithNext=True, wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "DentalH2", fontName=bold_name, fontSize=13, leading=19, textColor=colors.HexColor(TEAL),
            spaceBefore=12, spaceAfter=6, keepWithNext=True, wordWrap="CJK",
        ),
        "h3": ParagraphStyle(
            "DentalH3", fontName=bold_name, fontSize=11, leading=17, textColor=colors.HexColor(INK),
            spaceBefore=9, spaceAfter=5, keepWithNext=True, wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "DentalBody", fontName=font_name, fontSize=9.5, leading=16, textColor=colors.HexColor(INK),
            spaceAfter=6, wordWrap="CJK",
        ),
        "quote": ParagraphStyle(
            "DentalQuote", fontName=font_name, fontSize=9.2, leading=15, textColor=colors.HexColor(INK),
            leftIndent=9, rightIndent=7, borderWidth=0.6, borderColor=colors.HexColor("#B8DCD6"),
            borderPadding=8, backColor=colors.HexColor(PALE_TEAL), spaceBefore=5, spaceAfter=9, wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "DentalBullet", fontName=font_name, fontSize=9.4, leading=15.5, textColor=colors.HexColor(INK),
            leftIndent=16, firstLineIndent=-8, bulletIndent=3, spaceAfter=4, wordWrap="CJK",
        ),
        "caption": ParagraphStyle(
            "DentalCaption", fontName=font_name, fontSize=8, leading=12, textColor=colors.HexColor(MUTED),
            alignment=TA_CENTER, spaceBefore=4, spaceAfter=10, wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "DentalTableHeader", fontName=bold_name, fontSize=7.7, leading=11, textColor=colors.HexColor(WHITE),
            alignment=TA_CENTER, wordWrap="CJK",
        ),
        "table_body": ParagraphStyle(
            "DentalTableBody", fontName=font_name, fontSize=7.6, leading=11.5, textColor=colors.HexColor(INK),
            alignment=TA_LEFT, wordWrap="CJK",
        ),
        "table_number": ParagraphStyle(
            "DentalTableNumber", fontName=font_name, fontSize=7.6, leading=11.5, textColor=colors.HexColor(INK),
            alignment=TA_RIGHT, wordWrap="CJK",
        ),
        "footer": ParagraphStyle(
            "DentalFooter", fontName=font_name, fontSize=7.5, leading=10, textColor=colors.HexColor(MUTED),
            alignment=TA_LEFT, wordWrap="CJK",
        ),
    }

    doc = SimpleDocTemplate(
        str(output), pagesize=A4, leftMargin=margin_x, rightMargin=margin_x,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=title, author="Dental AI Assistant", subject="Auxiliary dental-image recognition report",
    )

    report_type = _metadata_value(cover_rows, {"报告类型", "report type"}, "辅助识别报告" if not english else "Auxiliary Report")
    generated_at = _metadata_value(cover_rows, {"生成时间", "generated at"}, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    document_id = sha256(markdown.encode("utf-8")).hexdigest()[:12].upper()
    story: list[object] = [
        Spacer(1, 8 * mm),
        Paragraph("DENTAL AI · AUXILIARY RECOGNITION" if english else "DENTAL AI · 牙科影像辅助识别", styles["cover_kicker"]),
        Paragraph(_reportlab_inline(title), styles["cover_title"]),
        Paragraph(_reportlab_inline(report_type), styles["cover_subtitle"]),
        Spacer(1, 22 * mm),
    ]

    metadata = list(cover_rows)
    metadata.append(("Document ID" if english else "报告编号", document_id))
    cards: list[list[object]] = []
    for offset in range(0, len(metadata), 2):
        card_row: list[object] = []
        for key, value in metadata[offset : offset + 2]:
            card_row.append(
                [Paragraph(_reportlab_inline(key), styles["cover_label"]), Paragraph(_reportlab_inline(value), styles["cover_value"])]
            )
        if len(card_row) == 1:
            card_row.append("")
        cards.append(card_row)
    metadata_table = Table(cards, colWidths=[available_width / 2 - 3 * mm, available_width / 2 - 3 * mm], hAlign="LEFT")
    metadata_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(PALE)),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 11),
                ("RIGHTPADDING", (0, 0), (-1, -1), 11),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story.extend(
        [
            metadata_table,
            Spacer(1, 10 * mm),
            Paragraph(
                "This report summarizes model-generated candidate regions and must be reviewed with the original image and clinical context."
                if english
                else "本报告汇总模型生成的疑似区域，必须结合原始影像、临床信息与专业人员复核后使用。",
                styles["quote"],
            ),
            Spacer(1, 3 * mm),
            Paragraph(
                f"Generated {html_escape(generated_at)} · ID {document_id}" if english else f"生成时间 {html_escape(generated_at)} · 报告编号 {document_id}",
                styles["footer"],
            ),
            PageBreak(),
        ]
    )

    def add_table(block: MarkdownBlock) -> None:
        rows = block.rows
        if not rows:
            return
        columns = len(rows[0])
        min_width = 16 * mm if columns <= 4 else (12 * mm if columns <= 7 else 8.5 * mm)
        widths = _column_widths(rows, available_width, min_width)
        table_rows: list[list[Paragraph]] = []
        for row_index, row in enumerate(rows):
            rendered_row: list[Paragraph] = []
            for value in row:
                style = styles["table_header"] if row_index == 0 else (
                    styles["table_number"] if re.fullmatch(r"[-+]?\d+(?:\.\d+)?%?|[-–—]", _clean_inline_text(value)) else styles["table_body"]
                )
                rendered_row.append(Paragraph(_reportlab_inline(value or "-"), style))
            table_rows.append(rendered_row)
        table = LongTable(table_rows, colWidths=widths, repeatRows=1, hAlign="LEFT", splitByRow=1)
        commands: list[tuple] = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(NAVY)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(WHITE)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("BOX", (0, 0), (-1, -1), 0.65, colors.HexColor("#BFD0DC")),
            ("LINEBELOW", (0, 1), (-1, -1), 0.35, colors.HexColor(LINE)),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]
        for row_index, row in enumerate(rows[1:], 1):
            if row_index % 2 == 0:
                commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor(PALE)))
            for col_index, value in enumerate(row):
                plain = _clean_inline_text(value)
                if "强烈建议" in plain or "Strong manual" in plain:
                    commands.append(("BACKGROUND", (col_index, row_index), (col_index, row_index), colors.HexColor(PALE_RED)))
                elif "建议人工" in plain or "Manual review" in plain:
                    commands.append(("BACKGROUND", (col_index, row_index), (col_index, row_index), colors.HexColor(PALE_AMBER)))
                elif "可信度较高" in plain or "High confidence" in plain:
                    commands.append(("BACKGROUND", (col_index, row_index), (col_index, row_index), colors.HexColor(PALE_TEAL)))
        table.setStyle(TableStyle(commands))
        story.extend([table, Spacer(1, 4 * mm)])

    def pdf_image(path_value: Path, max_width: float, max_height: float) -> Image | None:
        if not path_value.exists():
            return None
        try:
            with PILImage.open(path_value) as image:
                width_px, height_px = image.size
            scale = min(max_width / max(1, width_px), max_height / max(1, height_px))
            return Image(str(path_value), width=max(1, width_px * scale), height=max(1, height_px * scale), useDPI=True)
        except Exception:
            return None

    index = 0
    while index < len(body_blocks):
        block = body_blocks[index]
        if block.kind == "heading":
            level = max(1, min(3, block.level - 1))
            story.append(Paragraph(_reportlab_inline(block.text), styles[f"h{level}"]))
        elif block.kind == "paragraph":
            story.append(Paragraph(_reportlab_inline(block.text), styles["body"]))
        elif block.kind == "quote":
            story.append(Paragraph(_reportlab_inline(block.text), styles["quote"]))
        elif block.kind == "rule":
            story.extend([Spacer(1, 2 * mm), HRFlowable(width="100%", thickness=0.8, color=colors.HexColor(LINE)), Spacer(1, 2 * mm)])
        elif block.kind == "list":
            for item_index, item in enumerate(block.items, 1):
                bullet = f"{item_index}." if block.ordered else "•"
                story.append(Paragraph(_reportlab_inline(item), styles["bullet"], bulletText=bullet))
            story.append(Spacer(1, 2 * mm))
        elif block.kind == "table":
            add_table(block)
        elif block.kind == "image":
            image_blocks = [block]
            while index + 1 < len(body_blocks) and body_blocks[index + 1].kind == "image" and len(image_blocks) < 2:
                index += 1
                image_blocks.append(body_blocks[index])
            if len(image_blocks) == 2:
                cell_width = (available_width - 5 * mm) / 2
                cells: list[list[object]] = []
                for image_block in image_blocks:
                    flowable = pdf_image(image_block.image_path or Path(), cell_width, 66 * mm)
                    if flowable is None:
                        cells.append([Paragraph(_reportlab_inline(image_block.alt or "Image unavailable"), styles["caption"])])
                    else:
                        cells.append([flowable, Paragraph(_reportlab_inline(image_block.alt), styles["caption"])])
                image_table = Table([cells], colWidths=[cell_width, cell_width], hAlign="CENTER")
                image_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 3), ("RIGHTPADDING", (0, 0), (-1, -1), 3)]))
                story.extend([image_table, Spacer(1, 3 * mm)])
            else:
                image_block = image_blocks[0]
                flowable = pdf_image(image_block.image_path or Path(), available_width, 92 * mm)
                if flowable is not None:
                    story.extend([flowable, Paragraph(_reportlab_inline(image_block.alt), styles["caption"])])
                else:
                    story.append(Paragraph(_reportlab_inline(f"Image unavailable: {image_block.alt}"), styles["quote"]))
        index += 1

    def draw_first_page(canvas, document) -> None:
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(NAVY_DARK))
        canvas.rect(0, page_height - 74 * mm, page_width, 74 * mm, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor(TEAL))
        canvas.rect(0, page_height - 74 * mm, 5 * mm, 74 * mm, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor(CYAN))
        canvas.circle(page_width - 24 * mm, page_height - 19 * mm, 17 * mm, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor(NAVY_DARK))
        canvas.circle(page_width - 24 * mm, page_height - 19 * mm, 11 * mm, fill=1, stroke=0)
        canvas.setFont(font_name, 7.5)
        canvas.setFillColor(colors.HexColor(MUTED))
        canvas.drawRightString(page_width - margin_x, 9 * mm, f"1  /  {document_id}")
        canvas.restoreState()

    def draw_later_pages(canvas, document) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(LINE))
        canvas.setLineWidth(0.6)
        canvas.line(margin_x, page_height - 11 * mm, page_width - margin_x, page_height - 11 * mm)
        canvas.setFont(bold_name, 8)
        canvas.setFillColor(colors.HexColor(NAVY))
        canvas.drawString(margin_x, page_height - 8.5 * mm, "DENTAL AI")
        canvas.setFont(font_name, 7.5)
        canvas.setFillColor(colors.HexColor(MUTED))
        canvas.drawRightString(page_width - margin_x, page_height - 8.5 * mm, report_type)
        canvas.line(margin_x, 12 * mm, page_width - margin_x, 12 * mm)
        canvas.setFont(font_name, 7)
        canvas.drawString(margin_x, 8.5 * mm, "Research and auxiliary recognition only" if english else "仅供科研展示与辅助识别，不作为临床诊断依据")
        canvas.drawRightString(page_width - margin_x, 8.5 * mm, f"{document.page}  ·  {document_id}")
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_first_page, onLaterPages=draw_later_pages)
    return str(output)


def _docx_set_run_font(run, name: str, size_pt: float, color: str = INK, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if bold is not None:
        run.bold = bold


def _docx_add_inline(paragraph, text: str, font_name: str, size_pt: float, color: str = INK, force_bold: bool = False) -> None:
    for value, bold, code in _inline_tokens(text):
        run = paragraph.add_run(value)
        _docx_set_run_font(run, font_name, size_pt, TEAL if code else color, force_bold or bold)


def _docx_shade(element, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill.lstrip("#"))


def _docx_cell_margins(table, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table_properties = table._tbl.tblPr
    existing = table_properties.find(qn("w:tblCellMar"))
    if existing is not None:
        table_properties.remove(existing)
    margins = OxmlElement("w:tblCellMar")
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    table_properties.append(margins)


def _docx_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total = sum(widths_dxa)
    table.autofit = False
    table_properties = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = table_properties.find(qn(tag))
        if existing is not None:
            table_properties.remove(existing)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(total))
    table_width.set(qn("w:type"), "dxa")
    table_properties.append(table_width)
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), str(indent_dxa))
    table_indent.set(qn("w:type"), "dxa")
    table_properties.append(table_indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_properties.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_properties = cell._tc.get_or_add_tcPr()
            tc_width = tc_properties.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_properties.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")


def _docx_repeat_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        properties.append(marker)
    marker.set(qn("w:val"), "true")


def _docx_keep_row_together(row) -> None:
    """Prevent a table row from being split into fragments across pages."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:cantSplit"))
    if marker is None:
        marker = OxmlElement("w:cantSplit")
        properties.append(marker)
    marker.set(qn("w:val"), "true")


def _docx_add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def export_docx_from_markdown(markdown: str, path: str | Path, base_dir: str | Path) -> str:
    """Create an editable A4 Word report with native headings, tables, and figures."""
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt, RGBColor

    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    blocks = parse_report_markdown(markdown, base_dir)
    title, cover_rows, body_blocks = _extract_cover(blocks)
    english = _is_english(cover_rows, title)
    font_name = "Microsoft YaHei" if not english else "Arial"
    document_id = sha256(markdown.encode("utf-8")).hexdigest()[:12].upper()
    report_type = _metadata_value(cover_rows, {"报告类型", "report type"}, "辅助识别报告" if not english else "Auxiliary Report")

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    section.different_first_page_header_footer = True

    document.core_properties.title = title
    document.core_properties.subject = report_type
    document.core_properties.author = "Dental AI Assistant"
    document.core_properties.keywords = "dental imaging, auxiliary recognition, review"
    document.core_properties.comments = f"Document ID: {document_id}"

    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 16, 8),
        ("Heading 2", 13.5, TEAL, 12, 6),
        ("Heading 3", 11.5, INK, 9, 5),
    ):
        style = document.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Cm(0.65)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    caption_style = document.styles["Caption"]
    caption_style.font.name = font_name
    caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    caption_style.font.size = Pt(8.5)
    caption_style.font.italic = False
    caption_style.font.color.rgb = RGBColor.from_string(MUTED.lstrip("#"))
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(3)
    caption_style.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _docx_add_inline(header, "DENTAL AI", font_name, 8, NAVY, True)
    header.add_run("    ")
    _docx_add_inline(header, report_type, font_name, 8, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _docx_add_inline(
        footer,
        "仅供科研展示与辅助识别  ·  " if not english else "Research and auxiliary recognition only  ·  ",
        font_name,
        7.5,
        MUTED,
    )
    _docx_add_page_number(footer)
    _docx_add_inline(footer, f"  ·  {document_id}", font_name, 7.5, MUTED)

    # Editorial-cover opening block (A4-localized standard business brief).
    cover_spacer = document.add_paragraph()
    cover_spacer.paragraph_format.space_after = Pt(18)
    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.style = "Table Grid"
    _docx_table_geometry(banner, [9600], 120)
    _docx_cell_margins(banner, top=360, start=360, bottom=340, end=360)
    _docx_keep_row_together(banner.rows[0])
    banner_cell = banner.cell(0, 0)
    _docx_shade(banner_cell._tc, NAVY_DARK)
    kicker = banner_cell.paragraphs[0]
    kicker.paragraph_format.space_after = Pt(8)
    _docx_add_inline(kicker, "DENTAL AI · AUXILIARY RECOGNITION" if english else "DENTAL AI · 牙科影像辅助识别", font_name, 10, "#7DE0D1", True)
    title_paragraph = banner_cell.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(8)
    _docx_add_inline(title_paragraph, title, font_name, 27, WHITE, True)
    subtitle = banner_cell.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(0)
    _docx_add_inline(subtitle, report_type, font_name, 13, "#D9F4F0")
    document.add_paragraph().paragraph_format.space_after = Pt(5)

    metadata = list(cover_rows)
    metadata.append(("Document ID" if english else "报告编号", document_id))
    cover_table = document.add_table(rows=1, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_table.style = "Table Grid"
    for key, value in metadata:
        row = cover_table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        label_p = row.cells[0].paragraphs[0]
        value_p = row.cells[1].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        value_p.paragraph_format.space_after = Pt(0)
        _docx_add_inline(label_p, key, font_name, 9, MUTED, True)
        _docx_add_inline(value_p, value, font_name, 10, INK, True if key.lower() in {"报告类型", "report type"} else False)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    first_row = cover_table.rows[0]
    cover_table._tbl.remove(first_row._tr)
    _docx_table_geometry(cover_table, [2350, 7250], 120)
    _docx_cell_margins(cover_table, top=120, start=150, bottom=120, end=150)
    for row_index, row in enumerate(cover_table.rows):
        if row_index % 2 == 0:
            for cell in row.cells:
                _docx_shade(cell._tc, PALE)

    document.add_paragraph()
    callout = document.add_paragraph()
    callout.paragraph_format.left_indent = Cm(0.35)
    callout.paragraph_format.right_indent = Cm(0.35)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2
    _docx_shade(callout._p, PALE_TEAL)
    _docx_add_inline(
        callout,
        "This report summarizes model-generated candidate regions and must be reviewed with the original image and clinical context."
        if english
        else "本报告汇总模型生成的疑似区域，必须结合原始影像、临床信息与专业人员复核后使用。",
        font_name,
        10,
        INK,
        True,
    )
    callout.add_run().add_break(WD_BREAK.PAGE)

    def add_docx_table(block: MarkdownBlock) -> None:
        rows = block.rows
        if not rows:
            return
        columns = len(rows[0])
        table = document.add_table(rows=len(rows), cols=columns)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        font_size = 8.5 if columns <= 5 else (7.6 if columns <= 7 else 7.0)
        widths = _column_widths(rows, 9600, 850 if columns <= 7 else 620)
        widths_dxa = [max(1, int(width)) for width in widths]
        widths_dxa[-1] += 9600 - sum(widths_dxa)
        _docx_table_geometry(table, widths_dxa, 120)
        _docx_cell_margins(table)
        _docx_repeat_header(table.rows[0])
        for table_row in table.rows:
            _docx_keep_row_together(table_row)
        for row_index, values in enumerate(rows):
            for col_index, value in enumerate(values):
                cell = table.cell(row_index, col_index)
                cell.text = ""
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                plain = _clean_inline_text(value or "-")
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _docx_shade(cell._tc, NAVY)
                    _docx_add_inline(paragraph, value or "-", font_name, font_size, WHITE, True)
                else:
                    if re.fullmatch(r"[-+]?\d+(?:\.\d+)?%?|[-–—]", plain):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if row_index % 2 == 0:
                        _docx_shade(cell._tc, PALE)
                    if "强烈建议" in plain or "Strong manual" in plain:
                        _docx_shade(cell._tc, PALE_RED)
                    elif "建议人工" in plain or "Manual review" in plain:
                        _docx_shade(cell._tc, PALE_AMBER)
                    elif "可信度较高" in plain or "High confidence" in plain:
                        _docx_shade(cell._tc, PALE_TEAL)
                    _docx_add_inline(paragraph, value or "-", font_name, font_size, INK)
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

    for block in body_blocks:
        if block.kind == "heading":
            level = max(1, min(3, block.level - 1))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _docx_add_inline(paragraph, block.text, font_name, {1: 17, 2: 13.5, 3: 11.5}[level], {1: NAVY, 2: TEAL, 3: INK}[level], True)
        elif block.kind == "paragraph":
            paragraph = document.add_paragraph()
            _docx_add_inline(paragraph, block.text, font_name, 10.5, INK)
        elif block.kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.right_indent = Cm(0.35)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.18
            _docx_shade(paragraph._p, PALE_TEAL)
            _docx_add_inline(paragraph, block.text, font_name, 9.8, INK)
        elif block.kind == "rule":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            properties = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "6")
            top.set(qn("w:color"), LINE.lstrip("#"))
            borders.append(top)
            properties.append(borders)
        elif block.kind == "list":
            style_name = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                paragraph = document.add_paragraph(style=style_name)
                _docx_add_inline(paragraph, item, font_name, 10, INK)
        elif block.kind == "table":
            add_docx_table(block)
        elif block.kind == "image":
            image_path = block.image_path or Path()
            if image_path.exists():
                try:
                    with PILImage.open(image_path) as image:
                        width_px, height_px = image.size
                    max_width_cm = 16.8
                    max_height_cm = 10.0
                    scale = min(max_width_cm / max(1, width_px), max_height_cm / max(1, height_px))
                    width_cm = max(1.0, width_px * scale)
                    height_cm = max(1.0, height_px * scale)
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.keep_with_next = True
                    shape = paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
                    shape._inline.docPr.set("descr", block.alt or image_path.name)
                    caption = document.add_paragraph(style="Caption")
                    _docx_add_inline(caption, block.alt, font_name, 8.5, MUTED)
                except Exception:
                    paragraph = document.add_paragraph()
                    _docx_add_inline(paragraph, f"图片暂不可用：{block.alt}", font_name, 9.5, MUTED)
            else:
                paragraph = document.add_paragraph()
                _docx_add_inline(paragraph, f"图片暂不可用：{block.alt}", font_name, 9.5, MUTED)

    document.save(output)
    return str(output)
